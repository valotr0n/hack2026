from __future__ import annotations

import asyncio
import csv
import inspect
import json
import logging
import re
from io import BytesIO, StringIO
from pathlib import Path
from time import perf_counter
from typing import Any, AsyncIterator, Literal
from uuid import uuid4

import httpx
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document
from openai import AsyncOpenAI
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

from .config import settings
from .schemas import ChatHistoryMessage

logger = logging.getLogger(__name__)

TOP_K = 5
EmbeddingPromptName = Literal["search_query", "search_document"]
EmbeddingBackend = Literal["local", "open"]

SYSTEM_PROMPT = """\
Ты — точный аналитик документов. Работаешь с материалами, загруженными пользователем.

ПРАВИЛА (строго обязательны):
1. Отвечай ТОЛЬКО на основе фрагментов ниже — не используй общие знания и не додумывай.
2. Если ответа нет ни в одном фрагменте — скажи: «В загруженных документах эта информация отсутствует».
3. При цитировании указывай источник в скобках: (Фрагмент N).
4. Если информация есть в нескольких фрагментах — синтезируй её в единый ответ.
5. Противоречия между фрагментами — явно укажи: «В документах есть расхождение: ...».
6. Отвечай на русском языке. Структурируй ответ: сначала суть, потом детали.

ФРАГМЕНТЫ ДОКУМЕНТОВ:
{context}"""


def _table_rows_to_text(headers: list[str], rows: list[list[str]]) -> str:
    """Конвертирует табличные данные в читаемый текст для RAG-индексации."""
    lines: list[str] = []
    if headers:
        lines.append("Колонки: " + ", ".join(str(h) for h in headers))
        lines.append("")
    for i, row in enumerate(rows, start=1):
        if not any(str(cell).strip() for cell in row):
            continue
        if headers:
            parts = [f"{str(h)}={str(v)}" for h, v in zip(headers, row) if str(v).strip()]
        else:
            parts = [str(v) for v in row if str(v).strip()]
        if parts:
            lines.append(f"Строка {i}: " + ", ".join(parts))
    return "\n".join(lines)


def _extract_text_from_csv(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            text_io = StringIO(payload.decode(encoding))
            break
        except UnicodeDecodeError:
            continue
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="CSV file encoding not supported. Use UTF-8 or Windows-1251.",
        )

    reader = csv.reader(text_io)
    all_rows = [row for row in reader if any(cell.strip() for cell in row)]
    if not all_rows:
        return ""
    headers = all_rows[0]
    rows = all_rows[1:]
    return _table_rows_to_text(headers, rows)


def _extract_text_from_xlsx(payload: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(payload), read_only=True, data_only=True)
    sheets_text: list[str] = []

    for sheet in wb.worksheets:
        all_rows = [
            [str(cell.value) if cell.value is not None else "" for cell in row]
            for row in sheet.iter_rows()
            if any(cell.value is not None for cell in row)
        ]
        if not all_rows:
            continue
        headers = all_rows[0]
        rows = all_rows[1:]
        sheet_text = f"Лист: {sheet.title}\n{_table_rows_to_text(headers, rows)}"
        sheets_text.append(sheet_text)

    wb.close()
    return "\n\n".join(sheets_text)


def _normalize_pdf_text(text: str) -> str:
    """Убирает типовые артефакты browser-to-PDF печати."""
    cleaned_lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue

        if line.startswith(("file://", "http://", "https://")):
            continue

        if re.fullmatch(r"\d+/\d+", line):
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def build_upload_preview(full_text: str, limit: int = 500) -> str:
    """Делает preview так, чтобы визуальный блок не терялся за длинным текстом."""
    normalized = full_text.strip()
    if len(normalized) <= limit:
        return normalized

    visual_marker = "--- Визуальные элементы документа ---"
    if visual_marker not in normalized:
        return normalized[:limit]

    text_part, visual_part = normalized.split(visual_marker, 1)
    visual_preview = f"{visual_marker}{visual_part}".strip()
    if len(visual_preview) >= limit:
        return visual_preview[:limit]

    remaining = limit - len(visual_preview)
    text_preview = text_part.strip()[:remaining].strip()
    if not text_preview:
        return visual_preview

    return f"{visual_preview}\n\n{text_preview}".strip()[:limit]


async def extract_text_from_upload(file: UploadFile) -> str:
    from .vision import extract_image_descriptions

    filename = file.filename or "document"
    suffix = Path(filename).suffix.lower()
    payload = await file.read()
    started_at = perf_counter()
    vision_elapsed = 0.0

    text_extract_started_at = perf_counter()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(payload))
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages)
        text = _normalize_pdf_text(text)
    elif suffix == ".docx":
        document = DocxDocument(BytesIO(payload))
        text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs)
    elif suffix == ".txt":
        try:
            text = payload.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="TXT files must be UTF-8 encoded.",
            ) from exc
    elif suffix == ".csv":
        text = _extract_text_from_csv(payload)
    elif suffix in (".xlsx", ".xls"):
        text = _extract_text_from_xlsx(payload)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type. Use PDF, DOCX, TXT, CSV, or XLSX.",
        )
    text_extract_elapsed = perf_counter() - text_extract_started_at

    # Склеиваем переносы слов через дефис (артефакт PDF-парсинга)
    # "зави-\nсимости" → "зависимости"
    dehyphenated = re.sub(r"-\n(\S)", r"\1", text)
    normalized_text = "\n".join(line for line in dehyphenated.splitlines() if line.strip()).strip()

    if settings.vision_enabled and suffix in (".pdf", ".docx"):
        effective_max_images_per_document = settings.vision_max_images_per_document
        effective_max_new_tokens = settings.vision_max_new_tokens
        if len(payload) >= settings.vision_large_document_bytes_threshold:
            effective_max_images_per_document = min(
                effective_max_images_per_document,
                settings.vision_large_document_max_images,
            )
            effective_max_new_tokens = min(
                effective_max_new_tokens,
                settings.vision_large_document_max_new_tokens,
            )
            logger.info(
                "Large document vision throttle enabled: filename=%s bytes=%d max_images=%d max_new_tokens=%d",
                filename,
                len(payload),
                effective_max_images_per_document,
                effective_max_new_tokens,
            )

        vision_started_at = perf_counter()
        image_descriptions = await extract_image_descriptions(
            payload,
            suffix,
            settings.vision_model_id,
            effective_max_new_tokens,
            settings.vision_max_image_side,
            settings.vision_min_image_side,
            effective_max_images_per_document,
        )
        vision_elapsed = perf_counter() - vision_started_at
        if image_descriptions:
            visual_block = "\n\n".join(image_descriptions)
            normalized_text = f"{normalized_text}\n\n--- Визуальные элементы документа ---\n{visual_block}"

    if not normalized_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file does not contain readable text.",
        )

    logger.info(
        "Extracted upload text: filename=%s suffix=%s chars=%d text_extract=%.2fs vision=%.2fs total=%.2fs",
        filename,
        suffix,
        len(normalized_text),
        text_extract_elapsed,
        vision_elapsed,
        perf_counter() - started_at,
    )

    return normalized_text


def split_text_into_chunks(full_text: str) -> list[str]:
    splitter = SentenceSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    nodes = splitter.get_nodes_from_documents([Document(text=full_text)])
    chunks = [content for node in nodes if (content := node.get_content().strip())]

    if not chunks:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unable to split the document into chunks.",
        )

    return chunks


def get_embedding_dimension(embedding_model: SentenceTransformer) -> int:
    dimension = embedding_model.get_sentence_embedding_dimension()
    if not isinstance(dimension, int) or dimension <= 0:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Embedding model returned an invalid embedding dimension.",
        )
    return dimension


def _encode_texts(
    embedding_model: SentenceTransformer,
    texts: list[str],
    prompt_name: EmbeddingPromptName | None,
) -> list[list[float]]:
    encode_kwargs: dict[str, Any] = {
        "convert_to_numpy": True,
        "normalize_embeddings": True,
        "show_progress_bar": False,
        "batch_size": settings.embedding_batch_size,
    }
    if prompt_name is not None:
        encode_kwargs["prompt_name"] = prompt_name

    embeddings = embedding_model.encode(texts, **encode_kwargs)
    return embeddings.tolist()


async def fetch_embeddings(
    embedding_model: SentenceTransformer,
    texts: list[str],
    prompt_name: EmbeddingPromptName | None = None,
) -> list[list[float]]:
    if not texts:
        return []

    started_at = perf_counter()
    try:
        embeddings = await asyncio.to_thread(
            _encode_texts,
            embedding_model,
            texts,
            prompt_name,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Local embedding generation failed: {exc}",
        ) from exc
    logger.info(
        "Embeddings generated: backend=local texts=%d batch_size=%d total=%.2fs",
        len(texts),
        settings.embedding_batch_size,
        perf_counter() - started_at,
    )

    if len(embeddings) != len(texts):
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Embedding model returned an unexpected number of vectors.",
        )

    vector_size = len(embeddings[0]) if embeddings else 0
    if vector_size <= 0:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Embedding backend returned an invalid vector size.",
        )
    for embedding in embeddings:
        if not isinstance(embedding, list) or len(embedding) != vector_size:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail="Embedding model returned an invalid vector size.",
            )

    return embeddings


async def _qdrant_request(
    method: str,
    path: str,
    body: dict[str, Any] | None = None,
) -> dict[str, Any]:
    url = f"{settings.qdrant_url.rstrip('/')}{path}"
    request_kwargs: dict[str, Any] = {}
    if body is not None:
        request_kwargs["json"] = body

    try:
        async with httpx.AsyncClient(timeout=30.0, trust_env=False) as client:
            response = await client.request(method, url, **request_kwargs)
            response.raise_for_status()
            return response.json()
    except httpx.HTTPStatusError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=(
                f"Qdrant request failed for {url} with status "
                f"{exc.response.status_code}: {exc.response.text!r}"
            ),
        ) from exc
    except httpx.RequestError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Qdrant request failed for {url}: {exc}",
        ) from exc


async def _create_qdrant_collection_via_rest(doc_id: str, embedding_dimension: int) -> None:
    await _qdrant_request(
        "PUT",
        f"/collections/{doc_id}",
        {"vectors": {"size": embedding_dimension, "distance": "Cosine"}},
    )


async def _collection_exists(collection_id: str) -> bool:
    url = f"{settings.qdrant_url.rstrip('/')}/collections/{collection_id}"
    try:
        async with httpx.AsyncClient(timeout=10.0, trust_env=False) as client:
            response = await client.get(url)
            return response.status_code == 200
    except httpx.RequestError:
        return False


async def _upsert_qdrant_points_via_rest(
    doc_id: str,
    chunks: list[str],
    embeddings: list[list[float]],
    filename: str,
    source_id: str,
) -> None:
    points = [
        {
            "id": str(uuid4()),
            "vector": embedding,
            "payload": {
                "text": chunk,
                "source": filename,
                "chunk_index": index,
                "source_id": source_id,
                "embedding_backend": "local",
            },
        }
        for index, (chunk, embedding) in enumerate(zip(chunks, embeddings))
    ]
    await _qdrant_request(
        "PUT",
        f"/collections/{doc_id}/points?wait=true",
        {"points": points},
    )


async def create_document_collection(
    embedding_dimension: int,
    filename: str,
    chunks: list[str],
    embeddings: list[list[float]],
    notebook_id: str | None = None,
    source_id: str | None = None,
) -> str:
    collection_id = notebook_id if notebook_id else str(uuid4())
    sid = source_id if source_id else str(uuid4())

    if not await _collection_exists(collection_id):
        await _create_qdrant_collection_via_rest(collection_id, embedding_dimension)
    else:
        existing_backend = await get_collection_embedding_backend(collection_id)
        if existing_backend == "open":
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=(
                    "Блокнот был проиндексирован внешним эмбеддером. "
                    "Очистите источники и загрузите их заново после перехода на локальный режим."
                ),
            )
        current_vector_size = await get_collection_vector_size(collection_id)
        if current_vector_size is not None and current_vector_size != embedding_dimension:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=(
                    "Коллекция использует несовместимый размер векторов. "
                    "Очистите блокнот и загрузите источники заново."
                ),
            )

    await _upsert_qdrant_points_via_rest(
        collection_id,
        chunks,
        embeddings,
        filename,
        sid,
    )
    return collection_id


async def get_collection_vector_size(collection_id: str) -> int | None:
    if not await _collection_exists(collection_id):
        return None

    result = await _qdrant_request("GET", f"/collections/{collection_id}")
    collection_result = result.get("result") or {}
    vectors = (
        collection_result.get("config", {})
        .get("params", {})
        .get("vectors")
    )
    if isinstance(vectors, dict):
        size = vectors.get("size")
        if isinstance(size, int) and size > 0:
            return size
        default_vector = next(iter(vectors.values()), None)
        if isinstance(default_vector, dict):
            nested_size = default_vector.get("size")
            if isinstance(nested_size, int) and nested_size > 0:
                return nested_size
    return None


async def get_collection_embedding_backend(collection_id: str) -> EmbeddingBackend | None:
    if not await _collection_exists(collection_id):
        return None

    result = await _qdrant_request(
        "POST",
        f"/collections/{collection_id}/points/scroll",
        {
            "limit": 1,
            "with_payload": True,
            "with_vector": False,
        },
    )
    points = (result.get("result") or {}).get("points") or []
    if not points:
        return None

    payload = points[0].get("payload") or {}
    return "open" if payload.get("embedding_backend") == "open" else "local"


async def get_source_content(
    notebook_id: str,
    source_id: str,
    filename: str | None = None,
) -> dict[str, Any]:
    """Возвращает текст чанков конкретного source_id внутри коллекции ноутбука."""
    async def _scroll_points(filter_body: dict[str, Any]) -> list[dict[str, Any]]:
        points: list[dict[str, Any]] = []
        offset: Any = None

        while True:
            body: dict[str, Any] = {
                "limit": 1000,
                "with_payload": True,
                "with_vector": False,
                "filter": filter_body,
            }
            if offset is not None:
                body["offset"] = offset

            result = await _qdrant_request(
                "POST",
                f"/collections/{notebook_id}/points/scroll",
                body,
            )
            result_data = result.get("result") or {}
            points.extend(result_data.get("points") or [])
            offset = result_data.get("next_page_offset")
            if offset is None:
                break

        return points

    source_filter = {"must": [{"key": "source_id", "match": {"value": source_id}}]}
    all_points = await _scroll_points(source_filter)

    if not all_points and filename:
        filename_filter = {"must": [{"key": "source", "match": {"value": filename}}]}
        all_points = await _scroll_points(filename_filter)

    all_points.sort(key=lambda p: p.get("payload", {}).get("chunk_index", 0))
    texts = [
        p["payload"]["text"]
        for p in all_points
        if p.get("payload", {}).get("text", "").strip()
    ]
    filename = all_points[0]["payload"].get("source", "") if all_points else ""
    return {"text": "\n\n".join(texts), "filename": filename}


async def delete_source_chunks(collection_id: str, source_id: str) -> None:
    await _qdrant_request(
        "POST",
        f"/collections/{collection_id}/points/delete",
        {"filter": {"must": [{"key": "source_id", "match": {"value": source_id}}]}},
    )


async def delete_collection(collection_id: str) -> None:
    try:
        await _qdrant_request("DELETE", f"/collections/{collection_id}")
    except HTTPException:
        pass


async def get_notebook_content(collection_id: str) -> dict[str, Any]:
    all_points: list[dict[str, Any]] = []
    offset: Any = None

    while True:
        body: dict[str, Any] = {"limit": 1000, "with_payload": True, "with_vector": False}
        if offset is not None:
            body["offset"] = offset

        result = await _qdrant_request(
            "POST",
            f"/collections/{collection_id}/points/scroll",
            body,
        )
        result_data = result.get("result") or {}
        all_points.extend(result_data.get("points") or [])
        offset = result_data.get("next_page_offset")
        if offset is None:
            break

    all_points.sort(key=lambda p: (
        p.get("payload", {}).get("source", ""),
        p.get("payload", {}).get("chunk_index", 0),
    ))

    texts = [
        p["payload"]["text"]
        for p in all_points
        if p.get("payload", {}).get("text", "").strip()
    ]
    sources = sorted({
        p["payload"]["source"]
        for p in all_points
        if p.get("payload", {}).get("source")
    })

    return {"text": "\n\n".join(texts), "sources": sources}


async def search_document_chunks(
    doc_id: str,
    query_embedding: list[float],
) -> list[dict[str, Any]]:
    await _qdrant_request("GET", f"/collections/{doc_id}")
    query_result = await _qdrant_request(
        "POST",
        f"/collections/{doc_id}/points/query",
        {
            "query": query_embedding,
            "limit": TOP_K,
            "with_payload": True,
        },
    )

    result = query_result.get("result") or {}
    points = result.get("points") or []

    chunks: list[dict[str, Any]] = []
    for point in points:
        payload = point.get("payload") or {}
        text = payload.get("text")
        if isinstance(text, str) and text.strip():
            chunks.append(
                {
                    "text": text,
                    "source": payload.get("source", "unknown"),
                    "chunk_index": payload.get("chunk_index"),
                }
            )

    if not chunks:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No relevant chunks were found for the requested document.",
        )

    return chunks


def build_system_prompt(chunks: list[dict[str, Any]]) -> tuple[str, list[str]]:
    sources = [chunk["text"] for chunk in chunks]
    context_parts = []
    for index, chunk in enumerate(chunks, start=1):
        context_parts.append(
            f"--- Фрагмент {index} | {chunk['source']} ---\n{chunk['text']}\n---"
        )

    context = "\n\n".join(context_parts)
    prompt = SYSTEM_PROMPT.format(context=context)
    return prompt, sources


_PRE_ANSWER_INSTRUCTION = (
    "Перед ответом мысленно:\n"
    "1. Найди все фрагменты, относящиеся к вопросу.\n"
    "2. Проверь — нет ли противоречий между ними.\n"
    "3. Сформулируй ответ кратко и точно.\n\n"
    "Вопрос: {query}"
)


def build_chat_messages(
    system_prompt: str,
    history: list[ChatHistoryMessage],
    query: str,
) -> list[dict[str, str]]:
    messages: list[dict[str, str]] = [{"role": "system", "content": system_prompt}]
    for item in history:
        if item.role in {"assistant", "user"} and item.content.strip():
            messages.append({"role": item.role, "content": item.content})
    messages.append({"role": "user", "content": _PRE_ANSWER_INSTRUCTION.format(query=query)})
    return messages


async def create_llm_stream(
    llm_client: AsyncOpenAI,
    messages: list[dict[str, str]],
):
    try:
        return await llm_client.chat.completions.create(
            model=settings.llm_model,
            messages=messages,
            stream=True,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"LLM API request failed: {exc}",
        ) from exc


async def stream_chat_response(stream: Any, sources: list[str]) -> AsyncIterator[str]:
    first = True
    try:
        async for chunk in stream:
            choices = getattr(chunk, "choices", None) or []
            if not choices:
                continue

            delta = choices[0].delta.content or ""
            if not delta:
                continue

            if first:
                payload = json.dumps({"delta": delta, "sources": sources}, ensure_ascii=False)
                first = False
            else:
                payload = json.dumps({"delta": delta}, ensure_ascii=False)
            yield f"data: {payload}\n\n"
    except Exception as exc:
        error_payload = json.dumps(
            {"delta": "", "sources": sources, "error": f"LLM stream failed: {exc}"},
            ensure_ascii=False,
        )
        yield f"data: {error_payload}\n\n"
    finally:
        close = getattr(stream, "close", None)
        if callable(close):
            result = close()
            if inspect.isawaitable(result):
                await result

    yield "data: [DONE]\n\n"
